"""
demo_02 文档解析引擎 — 核心逻辑
═══════════════════════════════════════════════════════════════
从 Swagger/OpenAPI 文档提取接口，输出标准化接口列表。

源码参考: backend/app/services/enhanced_document_parser.py (1286行)

核心设计:
  策略模式——parser_map 注册表，10+种格式各自独立解析，新增格式不影响已有代码。
  两层解析——标准解析(Swagger/paths遍历) + LLM辅助(Word/PDF非结构化文本) → 降级到规则引擎。
  归一化输出——所有路径输出相同 APIInterface schema，下游7个 demo 统一消费。
═══════════════════════════════════════════════════════════════
"""
import sys
import re
from pathlib import Path
from typing import List, Dict, Callable, Optional

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
from demo_common import add_project_to_sys_path, llm_client
add_project_to_sys_path(__file__)


# ══════════════════════════════════════════════════════════════
# 策略模式：解析器注册表 (backend/enhanced_document_parser.py L56-80 同款)
# ══════════════════════════════════════════════════════════════
parser_map: Dict[str, Callable] = {}  # 运行时注册：swagger/postman/apifox/jmx/...


# ══════════════════════════════════════════════════════════════
# 格式检测 —— 根据输入数据结构推断文档格式
# ══════════════════════════════════════════════════════════════

def detect_format(data) -> str:
    """
    推断文档格式类型。检测顺序从最明确特征到最通用特征。

    为什么要先检测格式？
      - 不同格式的字段名完全不同 (Postman 叫 request.url.raw, Swagger 叫 paths)
      - 先检测格式再做针对性解析，比"全部试一遍"高效
      - 未知格式统一走 LLM 解析路径
    """
    # 字符串类型特殊处理 — 必须在 isinstance(dict) 之前检测
    if isinstance(data, str):
        if data.strip().startswith(('curl ', 'curl\n')):
            return "curl"
        # JMX XML 原始字符串
        if 'jmeterTestPlan' in data or 'TestPlan' in data:
            return "jmx"
        if 'hashTree' in data and 'HTTPSamplerProxy' in data:
            return "jmx"
        return "unknown"

    if not isinstance(data, dict):
        return "unknown"

    # ① Postman Collection v2.x: 有 info.schema 且含 "postman"
    info = data.get('info', {})
    schema_url = str(info.get('schema', ''))
    if 'postman' in schema_url.lower() or 'getpostman' in schema_url.lower():
        return "postman"
    # Postman v2.0 也可能没有 schema，通过 item[0].request 检测
    if 'item' in data and isinstance(data.get('item'), list):
        first = data['item'][0] if data['item'] else {}
        if isinstance(first, dict) and 'request' in first:
            return "postman"

    # ② Apifox: 有 apiCollection 或 apifoxProject 字段
    if 'apiCollection' in data or 'apifoxProject' in data:
        return "apifox"
    if 'collections' in data and any(
        isinstance(c, dict) and ('apis' in c or 'items' in c)
        for c in data.get('collections', [])
    ):
        return "apifox"

    # ③ JMeter JMX: XML/HTML 结构，含 jmeterTestPlan 或 hashTree
    raw_str = str(data)
    if 'jmeterTestPlan' in raw_str or 'TestPlan' in raw_str:
        return "jmx"
    if 'hashTree' in raw_str and 'HTTPSamplerProxy' in raw_str:
        return "jmx"

    # ④ HAR (HTTP Archive): 有 log.entries
    log = data.get('log', {})
    if 'entries' in log and isinstance(log.get('entries'), list):
        first_entry = log['entries'][0] if log['entries'] else {}
        if isinstance(first_entry, dict) and 'request' in first_entry:
            return "har"

    # ⑤ CURL: 已在函数开头检测，此处仅用于 dict 被误包装的 CURL 场景（极少数）

    # ⑥ Swagger/OpenAPI: 有 openapi 或 swagger 版本号
    if 'openapi' in data or 'swagger' in data:
        return "swagger"

    # ⑦ 自定义 interfaces 数组
    if 'interfaces' in data:
        return "custom_interfaces"

    # ⑧ 有 paths 对象（可能是简化版 Swagger）
    if 'paths' in data:
        return "swagger"

    return "unknown"


# ══════════════════════════════════════════════════════════════
# 格式感知入口 —— 自动检测并路由到对应解析器
# ══════════════════════════════════════════════════════════════

def parse_document(data, filename: str = "") -> list:
    """
    格式感知的文档解析入口。
    先检测格式 → 路由到对应解析器 → 归一化输出。

    与 parse_swagger_document() 的区别：
      - parse_swagger_document() 只处理 Swagger/interfaces 两种格式（向后兼容）
      - parse_document() 支持全部 10 种格式 + 自动检测 + parser_map 扩展
    """
    fmt = detect_format(data)

    # ── 路由1: parser_map 注册的解析器 ──
    if fmt in parser_map:
        raw = parser_map[fmt](data)
        return [_normalize_interface(r) for r in raw]

    # ── 路由2: Swagger/OpenAPI 标准格式 ──
    if fmt == "swagger":
        return parse_swagger_document(data)

    # ── 路由3: 自定义 interfaces 数组 ──
    if fmt == "custom_interfaces":
        return parse_swagger_document(data)

    # ── 路由4: JMX (XML格式, 需要结构化预处理) ──
    if fmt == "jmx":
        raw = _parse_jmx(data)
        return [_normalize_interface(r) for r in raw]

    # ── 路由5: CURL (纯文本, 正则提取) ──
    if fmt == "curl":
        raw = _parse_curl(str(data))
        return [_normalize_interface(r) for r in raw]

    # ── 路由6: LLM 兜底 ──
    if _llm_available():
        llm_result = _parse_via_llm(data)
        if llm_result:
            return [_normalize_interface(r) for r in llm_result]

    # ── 最终降级 ──
    return parse_swagger_document(data)

def parse_swagger_document(swagger_json: dict) -> list:
    """
    解析 Swagger/OpenAPI 文档，返回标准 APIInterface 列表。

    路由逻辑（backend L146 同款）:
      1. 有 'interfaces' 数组 → 自定义格式，直接归一化
      2. 有 'paths' 对象 → 标准 Swagger，遍历 paths 提取
      3. 未知格式 → LLM 辅助解析 → 降级到规则引擎

    Returns: List[Dict]，每个接口含 name/method/path/service/description/
             headers/params/body/response_schema/crud_type/category/version
    """
    interfaces = []
    parser_type = "rule_based_fallback"

    # 路由1: 自定义 interfaces 数组（demo 最常见格式）
    if 'interfaces' in swagger_json:
        raw_list = swagger_json['interfaces']
        interfaces = [_normalize_interface(raw) for raw in raw_list]
        parser_type = "swagger_direct"

        # LLM 增强：验证并补充缺失字段
        if _llm_available():
            interfaces = _llm_enhance(interfaces)
            parser_type = "llm_enhanced"

    # 路由2: 标准 Swagger paths（backend L178-250 同款）
    elif 'paths' in swagger_json:
        interfaces = _parse_swagger_paths(swagger_json)
        parser_type = "swagger_direct"

    # 路由3: LLM 解析未知格式
    elif _llm_available():
        llm_result = _parse_via_llm(swagger_json)
        if llm_result:
            interfaces = llm_result
            parser_type = "llm_enhanced"

    # 最终降级：当接口列表是原始 list 或单个 dict
    if not interfaces:
        src = swagger_json if isinstance(swagger_json, list) else [swagger_json]
        interfaces = [_normalize_interface(item) for item in src if isinstance(item, dict)]

    for iface in interfaces:
        iface['_parser_type'] = parser_type

    return interfaces


# ══════════════════════════════════════════════════════════════
# Swagger paths 遍历解析 (backend L178-250 同款)
# ══════════════════════════════════════════════════════════════

def _parse_swagger_paths(swagger_data: dict) -> list:
    """从标准 Swagger paths 对象提取接口列表"""
    interfaces = []
    paths = swagger_data.get('paths', {})
    servers = swagger_data.get('servers', [])
    base_url = servers[0].get('url', '') if servers else ''

    # 收集 definitions/components 用于 $ref 解析
    definitions = swagger_data.get('definitions', {})
    components = swagger_data.get('components', {})
    if 'schemas' in components:
        definitions = components['schemas']

    for path, path_item in paths.items():
        if not isinstance(path_item, dict):
            continue

        # path-level 共享 parameters
        shared_params = path_item.get('parameters', [])

        # 遍历 HTTP 方法
        for method in ['get', 'post', 'put', 'delete', 'patch']:
            op = path_item.get(method)
            if not isinstance(op, dict):
                continue

            # 合并 path-level + operation-level 参数
            all_params = shared_params + op.get('parameters', [])

            # 提取请求体 (OpenAPI 3.0 requestBody)
            body = {}
            req_body = op.get('requestBody', {})
            if isinstance(req_body, dict):
                body = _extract_content_schema(req_body.get('content', {}), definitions)

            # 提取响应 schema (200/201)
            response_schema = _extract_content_schema(
                op.get('responses', {}).get('200', {}).get('content', {}), definitions
            ) or _extract_content_schema(
                op.get('responses', {}).get('201', {}).get('content', {}), definitions
            )

            interfaces.append({
                "name": op.get('operationId') or op.get('summary') or f"{method.upper()} {path}",
                "method": method.upper(),
                "url": f"{base_url}{path}",
                "path": path,
                "service": _infer_service_from_path(path),
                "description": op.get('description', ''),
                "headers": _extract_headers(all_params),
                "params": _extract_query_params(all_params),
                "body": body,
                "response_schema": response_schema,
                "tags": op.get('tags', []),
                "deprecated": op.get('deprecated', False),
                "version": _extract_version(path),
            })

    return interfaces


def _extract_content_schema(content: dict, definitions: dict) -> dict:
    """从 content dict 中提取 schema（处理 application/json 等 content-type）"""
    for ct, ct_def in content.items():
        schema = ct_def.get('schema', {})
        if isinstance(schema, dict) and '$ref' in schema:
            schema = _resolve_ref(schema['$ref'], definitions)
        return {'content_type': ct, 'schema': schema}
    return {}


def _extract_headers(params: list) -> dict:
    """从参数列表提取 header 参数"""
    return {p['name']: p.get('description', '')
            for p in params if isinstance(p, dict) and p.get('in') == 'header'}


def _extract_query_params(params: list) -> dict:
    """从参数列表提取 query/path 参数"""
    result = {}
    for p in params:
        if not isinstance(p, dict) or p.get('in') not in ('query', 'path'):
            continue
        name = p.get('name', '')
        if name:
            result[name] = {
                "type": p.get('type', p.get('schema', {}).get('type', 'string')),
                "description": p.get('description', ''),
                "required": p.get('required', False),
            }
    return result


# ══════════════════════════════════════════════════════════════
# $ref 引用解析 (backend/document_parser.py L93-120 同款)
# ══════════════════════════════════════════════════════════════

def _resolve_ref(ref_path: str, definitions: dict, depth: int = 0) -> dict:
    """
    解析 JSON $ref 引用，如 "#/definitions/Device" → 展开后的 schema。
    为什么需要？Swagger 大量用 $ref 避免重复定义，不解引用就是空壳。
    深度限制防止循环引用 (A→B→A)。
    """
    if depth > 10:
        return {}
    parts = ref_path.replace('#/', '').split('/')
    result = definitions
    for part in parts:
        if isinstance(result, dict) and part in result:
            result = result[part]
        else:
            return {}
    if isinstance(result, dict) and '$ref' in result:
        return _resolve_ref(result['$ref'], definitions, depth + 1)
    return result if isinstance(result, dict) else {}


# ══════════════════════════════════════════════════════════════
# 接口归一化 (backend L312-380 的字段映射表同款)
# ══════════════════════════════════════════════════════════════

def _normalize_interface(raw: dict) -> dict:
    """
    将任意格式的接口数据归一化为标准 APIInterface schema。
    不同来源字段名不同（Swagger叫summary，Postman叫name），归一化消除差异。
    """
    url = raw.get('url', '')
    path = raw.get('path', url)
    method = str(raw.get('method', 'GET')).upper()

    return {
        "name": raw.get('name') or raw.get('summary') or raw.get('title') or f"{method} {path}",
        "method": method,
        "url": url,
        "path": path,
        "service": raw.get('service', _infer_service_from_path(path)),
        "description": raw.get('description', raw.get('desc', '')),
        "headers": raw.get('headers', {}),
        "params": raw.get('params', raw.get('parameters', raw.get('query_params', {}))),
        "body": raw.get('body', raw.get('requestBody', raw.get('request_body', {}))),
        "response_schema": raw.get('response_schema', raw.get('responses', {})),
        "tags": raw.get('tags', []),
        "deprecated": raw.get('deprecated', False),
        # ── 计算字段 ──
        "crud_type": _infer_crud_type(method),
        "version": raw.get('version', _extract_version(url)),
        "category": _infer_category(raw.get('tags', []), path),
    }


# ══════════════════════════════════════════════════════════════
# 推断函数 (backend L395-480 同款)
# ══════════════════════════════════════════════════════════════

def _infer_crud_type(method: str) -> str:
    """
    HTTP 方法 → CRUD 类型。下游 demo_03 用它判断依赖方向：
    CREATE → READ → UPDATE → DELETE 决定了用例执行顺序。
    """
    return {"POST": "CREATE", "GET": "READ", "PUT": "UPDATE",
            "PATCH": "UPDATE", "DELETE": "DELETE"}.get(method.upper(), "READ")


def _infer_service_from_path(path: str) -> str:
    """
    从 path 推断服务名。如 /api/v1/device/create → device-service。
    跳过 api、版本号 v1/v2、路径参数 {xxx} 等非业务段。
    """
    if not path:
        return "unknown-service"
    parts = [p for p in path.strip("/").split("/") if p
             and p.lower() not in ('api',)
             and not p.startswith('{')
             and not re.match(r'^v\d', p.lower())]
    return f"{parts[0]}-service" if parts else "api-service"


def _extract_version(url: str) -> str:
    """正则提取 API 版本：/api/v1/...→v1, /api/V0.1/...→v0.1"""
    if not url:
        return "v1"
    m = re.search(r'[/](v\d+(?:\.\d+)?)[/]', url, re.IGNORECASE)
    return m.group(1).lower() if m else "v1"


def _infer_category(tags: list, path: str = "") -> str:
    """
    从 tags[0] 或 path 推断业务分组。
    demo_03 用它按分组构建 CRUD 顺序链。
    """
    tag_map = {"auth": "account", "login": "account", "user": "account",
               "device": "device", "equipment": "device", "course": "course",
               "family": "family", "group": "family"}
    if tags:
        tag = tags[0].lower()
        return tag_map.get(tag, tag)
    if path:
        clean = re.sub(r'^(api/)?(v\d+(\.\d+)?/)?', '', path.strip("/").lower())
        seg = clean.split("/")[0]
        if seg and not seg.startswith("{"):
            return seg
    return "general"


# ══════════════════════════════════════════════════════════════
# LLM 辅助解析 (backend L500-600 的 AI 解析路径同款)
# ══════════════════════════════════════════════════════════════

def _llm_available() -> bool:
    """检查 LLM 是否可用（快速探测，不消耗 token）"""
    return llm_client._ensure_client()


def _parse_via_llm(doc_data: dict) -> Optional[List]:
    """
    LLM 从非标准文档提取接口信息。
    适用 Word/PDF/纯文本——这些没有结构化 schema，必须 AI 理解。
    LLM 不可用时返回 None，由上层规则引擎兜底。
    """
    import json as j
    doc_str = j.dumps(doc_data, ensure_ascii=False, indent=2)
    if len(doc_str) > 8000:
        doc_str = doc_str[:8000] + "\n...[文档已截断]..."

    prompt = f"""你是一位专业的API文档解析专家。请从以下文档内容中提取标准化的API接口信息。
要求输出JSON格式的接口列表，每个接口包含: name, method, path, description, headers, params, body, response_schema
请以纯JSON数组格式返回，不要包含markdown标记或其他额外文字。

文档内容:
{doc_str}
"""
    result = llm_client.extract_json(prompt, temperature=0.3, max_tokens=4000)
    if isinstance(result, dict) and 'interfaces' in result:
        return result['interfaces']
    if isinstance(result, list):
        return result
    return None


def _llm_enhance(interfaces: list) -> list:
    """
    LLM 增强：验证并修正规则引擎的结果。
    规则引擎做基础推断（GET→READ），LLM 检查是否有误（GET search 不是 READ）。
    这是"增强"而非"替换"——规则结果为基础，LLM 做修正。
    """
    import json as j
    brief = [{"index": i, "name": iface["name"], "method": iface["method"],
              "path": iface["path"], "crud_type": iface.get("crud_type", ""),
              "service": iface.get("service", "")} for i, iface in enumerate(interfaces)]

    prompt = f"""分析以下接口列表，验证 crud_type/service 是否正确。只修正有问题的字段。

原始列表:
{j.dumps(brief, ensure_ascii=False)}

请以JSON数组返回修正: [{{"index": 0, "crud_type": "READ"}}]
没问题的字段不要返回。"""
    result = llm_client.extract_json(prompt, temperature=0.2, max_tokens=2000)
    if isinstance(result, list):
        corrections = {item["index"]: item for item in result if "index" in item}
        for i, iface in enumerate(interfaces):
            if i in corrections:
                c = corrections[i]
                for field in ("crud_type", "service", "category"):
                    if field in c:
                        iface[field] = c[field]
    return interfaces


# ══════════════════════════════════════════════════════════════
# 格式解析器——Postman Collection v2.x (backend L650-720 同款)
# ══════════════════════════════════════════════════════════════

def _parse_postman(data: dict) -> list:
    """
    解析 Postman Collection v2.x 格式。
    字段映射: item[].request.method/url.raw/header/body → APIInterface。
    Postman 的 URL 是完整 URL (含 host)，需要从中提取 path。
    """
    interfaces = []
    items = data.get('item', [])

    def _walk(items, parent_tags=None):
        """递归遍历 Postman 的 folder/item 嵌套结构"""
        for item in items:
            if not isinstance(item, dict):
                continue
            # Folder 节点——递归进入，tags 叠加
            if 'item' in item and 'request' not in item:
                folder_name = item.get('name', '')
                sub_tags = (parent_tags or []) + [folder_name]
                _walk(item['item'], sub_tags)
                continue
            # API 节点
            req = item.get('request', {})
            if not isinstance(req, dict):
                continue
            url_info = req.get('url', {})
            url_raw = url_info.get('raw', '') if isinstance(url_info, dict) else str(url_info)
            path = url_info.get('path', []) if isinstance(url_info, dict) else []
            path_str = '/' + '/'.join(path) if path else url_raw

            headers = {}
            for h in req.get('header', []):
                if isinstance(h, dict):
                    headers[h.get('key', '')] = h.get('value', '')

            body = {}
            req_body = req.get('body', {})
            if isinstance(req_body, dict) and req_body.get('mode') == 'raw':
                body = {"content_type": "application/json", "raw": req_body.get('raw', '')}
            elif isinstance(req_body, dict) and req_body.get('mode') == 'formdata':
                body = {"content_type": "multipart/form-data",
                        "formdata": req_body.get('formdata', [])}

            interfaces.append({
                "name": item.get('name', f"{req.get('method','GET')} {path_str}"),
                "method": req.get('method', 'GET'),
                "url": url_raw,
                "path": path_str,
                "description": req.get('description', item.get('description', '')),
                "headers": headers,
                "params": _extract_postman_query(url_info),
                "body": body,
                "response_schema": _extract_postman_response(item),
                "tags": parent_tags or [],
            })
    _walk(items)
    return interfaces


def _extract_postman_query(url_info) -> dict:
    """从 Postman URL 对象提取 query 参数"""
    result = {}
    if not isinstance(url_info, dict):
        return result
    for q in url_info.get('query', []):
        if isinstance(q, dict):
            result[q.get('key', '')] = {
                "type": "string",
                "description": q.get('description', ''),
                "required": not q.get('disabled', False),
            }
    return result


def _extract_postman_response(item: dict) -> dict:
    """从 Postman item 提取示例响应 (作为 response_schema 参考)"""
    responses = item.get('response', [])
    if not responses:
        return {}
    first = responses[0] if isinstance(responses, list) and responses else {}
    if not isinstance(first, dict):
        return {}
    return {"status": first.get('code'), "body": first.get('body', '')}


# ══════════════════════════════════════════════════════════════
# 格式解析器——Apifox 导出格式 (backend L722-790 同款)
# ══════════════════════════════════════════════════════════════

def _parse_apifox(data: dict) -> list:
    """
    解析 Apifox 导出格式。支持两种结构:
      1. apiCollection 数组——Apifox 8.x 导出
      2. collections 数组——Apifox 9.x 导出
    每个 API 含: name/method/path/requestBody/parameters/responses
    """
    interfaces = []

    # 结构1: apiCollection 顶层数组
    api_list = data.get('apiCollection', [])
    # 结构2: collections[].apis 嵌套
    if not api_list:
        for col in data.get('collections', []):
            if isinstance(col, dict) and 'apis' in col:
                api_list.extend(col['apis'])

    for api in api_list:
        if not isinstance(api, dict):
            continue
        method = (api.get('method') or 'GET').upper()
        path = api.get('path', '')

        # 提取 headers
        headers = {}
        for h in api.get('headers', []):
            if isinstance(h, dict):
                headers[h.get('name', h.get('key', ''))] = h.get('value', '')

        # 提取 query params
        params = {}
        for p in api.get('parameters', api.get('query', [])):
            if isinstance(p, dict):
                name = p.get('name', p.get('key', ''))
                if name:
                    params[name] = {
                        "type": p.get('type', 'string'),
                        "description": p.get('description', ''),
                        "required": p.get('required', False),
                    }

        # 提取请求体
        body = {}
        req_body = api.get('requestBody', {})
        if isinstance(req_body, dict):
            body = {"content_type": "application/json", "schema": req_body}
            if 'raw' in req_body:
                body['raw'] = req_body['raw']

        # 提取响应 schema
        response_schema = {}
        resp_list = api.get('responses', [])
        if resp_list and isinstance(resp_list[0], dict):
            response_schema = resp_list[0]

        interfaces.append({
            "name": api.get('name', f"{method} {path}"),
            "method": method,
            "url": api.get('url', path),
            "path": path,
            "description": api.get('description', ''),
            "headers": headers,
            "params": params,
            "body": body,
            "response_schema": response_schema,
            "tags": api.get('tags', []),
        })
    return interfaces


# ══════════════════════════════════════════════════════════════
# 格式解析器——JMeter JMX (XML格式) (backend L792-860 同款)
# ══════════════════════════════════════════════════════════════

def _parse_jmx(data) -> list:
    """
    解析 JMeter JMX 格式 (XML)。使用 xml.etree 标准库。
    JMX 中 HTTPSamplerProxy 元素包含: testname, path, method, domain, port。
    为什么不用 lxml？减少外部依赖，stdlib 的 ElementTree 足够解析 JMX。
    """
    import xml.etree.ElementTree as ET

    interfaces = []
    raw_text = data if isinstance(data, str) else str(data)

    # JMX 可能是 dict 包装过的 XML 文本，或原始 XML 字符串
    if isinstance(data, dict):
        # 尝试把 dict 的字符串值拼接起来
        raw_text = _dict_to_text(data)

    try:
        root = ET.fromstring(raw_text)
    except ET.ParseError:
        # 尝试从原始 dict 的各字段中提取 XML 片段
        return _parse_jmx_fallback(data)

    # 遍历所有 HTTPSamplerProxy 元素
    for sampler in root.iter('HTTPSamplerProxy'):
        testname = sampler.get('testname', '')
        domain = sampler.findtext('stringProp[@name="HTTPSampler.domain"]', default='')
        port = sampler.findtext('stringProp[@name="HTTPSampler.port"]', default='80')
        path = sampler.findtext('stringProp[@name="HTTPSampler.path"]', default='')
        method = sampler.findtext('stringProp[@name="HTTPSampler.method"]', default='GET')
        protocol = sampler.findtext('stringProp[@name="HTTPSampler.protocol"]', default='https')

        url = f"{protocol}://{domain}:{port}{path}" if domain else path

        interfaces.append({
            "name": testname or f"{method} {path}",
            "method": method.upper(),
            "url": url,
            "path": path,
            "description": f"JMX 测试计划: {testname}",
            "headers": {},
            "params": {},
            "body": {},
            "response_schema": {},
            "tags": ["jmx"],
        })
    return interfaces


def _parse_jmx_fallback(data) -> list:
    """JMX 解析降级：从 dict 文本中正则提取接口信息"""
    interfaces = []
    raw_text = str(data)

    # 正则匹配 HTTPSamplerProxy 的属性
    pattern = r'testname="([^"]*)"[^>]*'
    matches = re.findall(pattern, raw_text)
    path_pattern = r'HTTPSampler\.path["\s>]+([^<]+)'
    method_pattern = r'HTTPSampler\.method["\s>]+([^<]+)'

    paths = re.findall(path_pattern, raw_text)
    methods = re.findall(method_pattern, raw_text)

    for i, name in enumerate(matches):
        path_val = paths[i] if i < len(paths) else '/'
        method_val = (methods[i] if i < len(methods) else 'GET').upper()
        interfaces.append({
            "name": name,
            "method": method_val,
            "url": path_val,
            "path": path_val,
            "description": f"JMX 接口: {name}",
            "headers": {}, "params": {}, "body": {}, "response_schema": {},
            "tags": ["jmx"],
        })
    return interfaces


def _dict_to_text(data: dict) -> str:
    """递归将 dict 转为纯文本（用于 JMX/XML 被包装成 JSON 的场景）"""
    parts = []
    for k, v in data.items():
        if isinstance(v, str):
            parts.append(v)
        elif isinstance(v, dict):
            parts.append(_dict_to_text(v))
        elif isinstance(v, list):
            for item in v:
                if isinstance(item, str):
                    parts.append(item)
                elif isinstance(item, dict):
                    parts.append(_dict_to_text(item))
    return '\n'.join(parts)


# ══════════════════════════════════════════════════════════════
# 格式解析器——CURL 命令 (backend L862-930 同款)
# ══════════════════════════════════════════════════════════════

def _parse_curl(curl_text: str) -> list:
    """
    解析 CURL 命令文本，提取单个接口信息。
    支持的格式:
      - URL: curl 'url' / curl "url" / curl url
      - Method: -X POST / --request PUT (默认 GET)
      - Header: -H 'Key: Value' / --header "Key: Value"
      - Body: -d 'data' / --data 'data' / --data-raw 'data'
    为什么用正则而非 shlex？CURL 参数格式多样，regex 更容错。
    """
    if not curl_text or not isinstance(curl_text, str):
        return []

    curl_text = curl_text.strip()

    # 提取 URL —— 支持单引号/双引号/裸URL
    url = ''
    url_match = re.search(r"""curl\s+(?:-[Xx]\s+\S+\s+)?(?:--request\s+\S+\s+)?['"]?((?:https?://|/)[^\s'"]+)""", curl_text)
    if not url_match:
        # 尝试匹配最后的 URL
        url_match = re.search(r"""['"]((?:https?://)[^'"]+)['"]""", curl_text)
    if url_match:
        url = url_match.group(1)

    # 提取 HTTP 方法
    method = 'GET'
    method_match = re.search(r'(?:-X|--request)\s+(\w+)', curl_text)
    if method_match:
        method = method_match.group(1).upper()
    elif re.search(r'(?:-d|--data)', curl_text):
        method = 'POST'  # 有 body 默认 POST

    # 提取 Headers —— 多种格式
    headers = {}
    for m in re.finditer(r"""(?:-H|--header)\s+['"]([^'"]+)['"]""", curl_text):
        parts = m.group(1).split(':', 1)
        if len(parts) == 2:
            headers[parts[0].strip()] = parts[1].strip()

    # 提取 Body
    body = {}
    body_match = re.search(r"""(?:--data(?:-raw|-binary|-urlencode)?|-d)\s+['"]([^'"]+)['"]""", curl_text)
    if body_match:
        body_raw = body_match.group(1)
        body = {"content_type": "application/json", "raw": body_raw}

    # 推断 path
    path = url
    if url.startswith('http'):
        path_match = re.search(r'https?://[^/]+(/[^\s?#]*)', url)
        if path_match:
            path = path_match.group(1)
    if not path:
        path = '/'

    return [{
        "name": f"{method} {path}",
        "method": method,
        "url": url,
        "path": path,
        "description": f"从 CURL 命令解析: {curl_text[:80]}...",
        "headers": headers,
        "params": {},
        "body": body,
        "response_schema": {},
        "tags": ["curl"],
    }]


# ══════════════════════════════════════════════════════════════
# 格式解析器——HAR (HTTP Archive) (backend L932-990 同款)
# ══════════════════════════════════════════════════════════════

def _parse_har(data: dict) -> list:
    """
    解析 HAR (HTTP Archive) 格式——浏览器抓包的 JSON 标准格式。
    字段映射: log.entries[].request.method/url/headers/postData → APIInterface。
    HAR 常用于从浏览器 DevTools 导出的请求集合中提取接口。
    """
    interfaces = []
    entries = data.get('log', {}).get('entries', [])

    for entry in entries:
        if not isinstance(entry, dict):
            continue
        req = entry.get('request', {})
        if not isinstance(req, dict):
            continue

        url_full = req.get('url', '')
        path = ''
        if url_full.startswith('http'):
            path_match = re.search(r'https?://[^/]+(/[^\s?#]*)', url_full)
            if path_match:
                path = path_match.group(1)
        if not path:
            path = url_full if url_full.startswith('/') else '/' + url_full

        # 提取 headers
        headers = {}
        for h in req.get('headers', []):
            if isinstance(h, dict):
                headers[h.get('name', '')] = h.get('value', '')

        # 提取 query 参数
        params = {}
        for q in req.get('queryString', []):
            if isinstance(q, dict):
                params[q.get('name', '')] = {
                    "type": "string",
                    "description": q.get('value', ''),
                    "required": False,
                }

        # 提取请求体
        body = {}
        post_data = req.get('postData', {})
        if isinstance(post_data, dict) and post_data.get('text'):
            body = {
                "content_type": post_data.get('mimeType', 'application/json'),
                "raw": post_data.get('text', ''),
            }

        # HAR response 信息可作为 response_schema 参考
        resp = entry.get('response', {})
        response_schema = {}
        if isinstance(resp, dict):
            resp_content = resp.get('content', {})
            if isinstance(resp_content, dict) and resp_content.get('text'):
                response_schema = {
                    "status": resp.get('status'),
                    "content_type": resp_content.get('mimeType'),
                    "body": resp_content.get('text', '')[:500],  # 截断大响应
                }

        interfaces.append({
            "name": f"{req.get('method', 'GET')} {path}",
            "method": req.get('method', 'GET').upper(),
            "url": url_full,
            "path": path,
            "description": f"HAR 请求: {url_full}",
            "headers": headers,
            "params": params,
            "body": body,
            "response_schema": response_schema,
            "tags": ["har"],
        })
    return interfaces


# ══════════════════════════════════════════════════════════════
# 格式解析器——PDF 文档解析 (可选依赖: PyPDF2)
# ══════════════════════════════════════════════════════════════

def _parse_pdf(filepath: str) -> list:
    """
    解析 PDF 文档中的接口描述 (可选依赖: PyPDF2)。
    PDF 是非结构化文档，提取文本后走 LLM 解析路径。
    PyPDF2 不可用时返回空列表（上层降级处理）。
    """
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        print("  [WARN] PyPDF2 未安装，无法解析 PDF。安装: pip install PyPDF2")
        return []

    try:
        reader = PdfReader(filepath)
        text = ''
        for page in reader.pages[:20]:  # 最多读 20 页
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
    except Exception as e:
        print(f"  [ERROR] PDF 解析失败: {e}")
        return []

    if not text.strip():
        return []

    # PDF 文本走 LLM 解析
    if _llm_available():
        return _parse_via_llm({"pdf_text": text}) or []

    # 无 LLM 时：正则粗略提取 URL 和 HTTP 方法
    interfaces = []
    url_pattern = r'(?:GET|POST|PUT|DELETE|PATCH)\s+(https?://[^\s]+)'
    for m in re.finditer(url_pattern, text, re.IGNORECASE):
        parts = m.group(0).split(maxsplit=1)
        interfaces.append({
            "name": f"{parts[0]} {parts[1][:50] if len(parts) > 1 else '/'}",
            "method": parts[0].upper(),
            "url": parts[1] if len(parts) > 1 else '/',
            "path": parts[1] if len(parts) > 1 else '/',
            "description": f"从 PDF 提取",
            "headers": {}, "params": {}, "body": {}, "response_schema": {},
            "tags": ["pdf"],
        })
    return interfaces


# ══════════════════════════════════════════════════════════════
# 格式解析器——Word 文档解析 (可选依赖: python-docx)
# ══════════════════════════════════════════════════════════════

def _parse_docx(filepath: str) -> list:
    """
    解析 Word (.docx) 文档中的接口描述 (可选依赖: python-docx)。
    提取段落和表格文本，走 LLM 解析路径。
    python-docx 不可用时返回空列表（上层降级处理）。
    """
    try:
        from docx import Document
    except ImportError:
        print("  [WARN] python-docx 未安装，无法解析 DOCX。安装: pip install python-docx")
        return []

    try:
        doc = Document(filepath)
        text_parts = []
        for para in doc.paragraphs[:200]:  # 最多 200 段落
            if para.text.strip():
                text_parts.append(para.text)

        # 也提取表格内容
        for table in doc.tables[:30]:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                text_parts.append(' | '.join(cells))
    except Exception as e:
        print(f"  [ERROR] DOCX 解析失败: {e}")
        return []

    text = '\n'.join(text_parts)
    if not text.strip():
        return []

    # DOCX 文本走 LLM 解析
    if _llm_available():
        return _parse_via_llm({"docx_text": text}) or []

    # 无 LLM 时正则粗略提取
    interfaces = []
    for m in re.finditer(r'(?:GET|POST|PUT|DELETE|PATCH)\s+(https?://[^\s]+)', text, re.IGNORECASE):
        parts = m.group(0).split(maxsplit=1)
        interfaces.append({
            "name": f"{parts[0]} {parts[1][:50] if len(parts) > 1 else '/'}",
            "method": parts[0].upper(),
            "url": parts[1] if len(parts) > 1 else '/',
            "path": parts[1] if len(parts) > 1 else '/',
            "description": f"从 Word 文档提取",
            "headers": {}, "params": {}, "body": {}, "response_schema": {},
            "tags": ["docx"],
        })
    return interfaces


# ══════════════════════════════════════════════════════════════
# 通用工具函数
# ══════════════════════════════════════════════════════════════

def _safe_json_parse(text: str) -> dict:
    """
    安全 JSON 解析：json.loads → ast.literal_eval → 空 dict。
    为什么需要？用户可能贴 Python dict 而非 JSON（单引号 vs 双引号），
    或者 LLM 返回的 JSON 带 markdown 标记、尾部逗号等。
    三层容错确保不崩溃。
    """
    import ast

    if not text or not isinstance(text, str):
        return {}

    # 去除 markdown 标记
    text = text.strip()
    if text.startswith('```'):
        lines = text.split('\n')
        text = '\n'.join(lines[1:-1] if lines[-1].strip() == '```' else lines[1:])

    # 层1: 标准 json.loads
    try:
        import json as j
        return j.loads(text)
    except (j.JSONDecodeError, ValueError):
        pass

    # 层2: ast.literal_eval (支持 Python 字面量)
    try:
        result = ast.literal_eval(text)
        return result if isinstance(result, dict) else {}
    except (ValueError, SyntaxError):
        pass

    return {}


def parse_file(filepath: str) -> list:
    """
    文件入口——自动检测扩展名 → 读取文件 → 解析 → 归一化输出。
    支持的扩展名: .json / .yaml / .txt / .pdf / .docx / .jmx / .har / .curl
    这是 parse_document() 的"文件版本"，暴露给 entry 脚本使用。
    """
    path = Path(filepath)
    if not path.exists():
        print(f"  [ERROR] 文件不存在: {filepath}")
        return []

    ext = path.suffix.lower()

    # PDF 和 DOCX 走专门的解析器
    if ext == '.pdf':
        raw = _parse_pdf(filepath)
        return [_normalize_interface(r) for r in raw]
    if ext == '.docx':
        raw = _parse_docx(filepath)
        return [_normalize_interface(r) for r in raw]

    # 文本类文件——读取内容后检测格式
    encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
    content = None
    for enc in encodings:
        try:
            content = path.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if content is None:
        print(f"  [ERROR] 无法解码文件: {filepath}")
        return []

    # JSON / YAML 文件
    if ext in ('.json', '.har'):
        import json as j
        try:
            data = j.loads(content)
        except j.JSONDecodeError:
            print(f"  [WARN] JSON 解析失败，尝试 LLM 解析")
            data = content
        return parse_document(data, path.name)

    if ext in ('.yaml', '.yml'):
        try:
            import yaml
            data = yaml.safe_load(content)
        except Exception:
            data = content
        return parse_document(data, path.name)

    # CURL 纯文本
    if ext == '.curl' or content.strip().startswith('curl '):
        return parse_document(content, path.name)

    # JMX XML
    if ext == '.jmx':
        return parse_document(content, path.name)

    # 其他文本类型——尝试 JSON → LLM
    try:
        import json as j
        data = j.loads(content)
    except j.JSONDecodeError:
        data = content
    return parse_document(data, path.name)


# ══════════════════════════════════════════════════════════════
# 解析器注册 (策略模式——模块末尾统一注册，新增格式只需加一行)
# ══════════════════════════════════════════════════════════════

parser_map["postman"] = _parse_postman
parser_map["apifox"] = _parse_apifox
parser_map["jmx"] = _parse_jmx
parser_map["curl"] = _parse_curl
parser_map["har"] = _parse_har
parser_map["pdf"] = _parse_pdf
parser_map["docx"] = _parse_docx
